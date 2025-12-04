#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
thesis_tools.reference_tools
============================

基于 Zotero 数据与 python-docx 的 Word 参考文献与正文引用工具。

设计目标：
- 提供可复用的“选文献 + 写参考文献列表 + 给句子追加 [n] 引用”的工具函数；
- 尽量复用已有的 optimize_proposal_references.py 中的模型与格式化逻辑；
- 让后续 Codex 会话可以快速发现并调用这些工具，而不必重复阅读脚本细节。
"""

from __future__ import annotations

from pathlib import Path
from typing import Dict, Iterable, List, Tuple
import re

from docx import Document

from .optimize_proposal_references import (
    ReferenceItem,
    format_reference_text,
    load_chinese_items_from_json,
    load_foreign_items_from_sqlite,
)


ROOT_DIR = Path(__file__).resolve().parent.parent


def sort_reference_key(ref: ReferenceItem) -> Tuple[int, str]:
    """用于参考文献排序的 key：按年份（新到旧）、标题排序。"""

    year = ref.year if ref.year is not None else 0
    title = ref.title or ""
    return year, title


def select_references_from_zotero(
    target_total: int = 30,
    min_foreign: int = 4,
    allow_test_items: bool = True,
) -> List[ReferenceItem]:
    """
    从本仓库 Zotero 数据源中选出一组参考文献条目。

    - 仅保留年份 <= 2025（或未填写年份）的条目；
    - 通过 min_foreign 控制外文文献的下限；
    - allow_test_items=True 时会保留类似 “测试文献” 这样的占位条目。
    """

    all_json_items = load_chinese_items_from_json()
    foreign_from_db = load_foreign_items_from_sqlite()

    def is_valid(ref: ReferenceItem) -> bool:
        title = (ref.title or "").strip().strip('"')
        if not title:
            return False
        if not allow_test_items and "测试文献" in title:
            return False
        if ref.year is not None and ref.year > 2025:
            return False
        return True

    # 外文文献：SQLite 为主
    foreign_items = [item for item in foreign_from_db if is_valid(item)]
    foreign_items.sort(key=sort_reference_key, reverse=True)

    # 中文文献：来自 JSON，排除 is_foreign
    chinese_items = [
        item
        for item in all_json_items
        if not item.is_foreign and is_valid(item)
    ]

    # 去重（按标题）
    unique_cn: Dict[str, ReferenceItem] = {}
    for item in chinese_items:
        title_key = (item.title or "").strip()
        if title_key and title_key not in unique_cn:
            unique_cn[title_key] = item
    chinese_list = sorted(unique_cn.values(), key=sort_reference_key, reverse=True)

    selected: List[ReferenceItem] = []

    if foreign_items:
        use_foreign = foreign_items[: max(min_foreign, len(foreign_items))]
        selected.extend(use_foreign)

    remaining = max(0, target_total - len(selected))
    if remaining > 0:
        selected.extend(chinese_list[:remaining])
    else:
        selected = selected[:target_total]

    if len(selected) < target_total:
        print(
            f"[WARN] 当前 Zotero 库中可用的参考文献仅 {len(selected)} 条，"
            f"未能达到目标数量 {target_total}。"
        )

    return selected


def delete_paragraph(paragraph) -> None:
    """安全删除一个段落对象（python-docx）。"""

    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def append_citation_suffix(base_text: str, indices: Iterable[int]) -> str:
    """
    在原句末尾追加形如 [1][2][3] 的引用标注。

    - 若原文本以句号“。”结尾，会在句号前插入引用；
    - 若末尾已有 [n] 形式的引用，会先清理掉旧的，再追加新的（去重后）；
    - indices 会去重并排序。
    """

    uniq_indices = sorted({int(i) for i in indices})
    if not uniq_indices:
        return base_text

    text = base_text.rstrip()
    has_period = text.endswith("。")
    if has_period:
        text = text[:-1].rstrip()

    # 移除末尾已有的 [数字]
    while True:
        m = re.search(r"\[\d+\]$", text)
        if not m:
            break
        text = text[: m.start()].rstrip()

    suffix = "".join(f"[{i}]" for i in uniq_indices)
    if has_period:
        return text + suffix + "。"
    return text + suffix


def replace_paragraph_text_in_row(
    doc: Document,
    row_index: int,
    para_index: int,
    new_text: str,
    table_index: int = 0,
) -> None:
    """
    在指定表格行的所有 5 列中，用统一的新文本替换目标段落内容。

    适用于“整行内容是同一模块”的表格型开题报告。
    """

    table = doc.tables[table_index]
    if row_index >= len(table.rows):
        raise IndexError(f"表格中不存在行索引 {row_index}。")

    row = table.rows[row_index]
    for cell in row.cells:
        if para_index >= len(cell.paragraphs):
            continue
        para = cell.paragraphs[para_index]
        para.text = new_text


def rewrite_reference_section_in_table(
    doc: Document,
    refs: List[ReferenceItem],
    row_index: int,
    heading_text: str = "五、主要参考文献（20篇以上）",
    table_index: int = 0,
) -> None:
    """
    在 Word 表格中重写“主要参考文献”列表。

    - 定位 table_index 对应表格的 row_index 行；
    - 保留每列第一个段落作为标题（若为空则创建），其余段落删除；
    - 在每个单元格中写入完整的 [1]...[N] 列表，保持 5 列内容一致。
    """

    if not doc.tables:
        raise ValueError("文档中未找到任何表格，无法定位参考文献区域。")

    table = doc.tables[table_index]
    if row_index >= len(table.rows):
        raise ValueError("表格行数不足，无法定位参考文献所在行。")

    refs_row = table.rows[row_index]

    first_cell = refs_row.cells[0]
    if not first_cell.paragraphs:
        heading_para = first_cell.add_paragraph(heading_text)
    else:
        heading_para = first_cell.paragraphs[0]
        if not heading_para.text.strip():
            heading_para.text = heading_text

    item_style = heading_para.style
    if len(first_cell.paragraphs) > 1:
        item_style = first_cell.paragraphs[1].style

    for cell in refs_row.cells:
        paras = list(cell.paragraphs)
        if not paras:
            heading = cell.add_paragraph(heading_text)
            heading.style = heading_para.style
            paras = [heading]

        for para in paras[1:]:
            delete_paragraph(para)

        for idx, ref in enumerate(refs, start=1):
            text = format_reference_text(ref, index=idx)
            p = cell.add_paragraph(text)
            p.style = item_style


__all__ = [
    "select_references_from_zotero",
    "sort_reference_key",
    "append_citation_suffix",
    "delete_paragraph",
    "replace_paragraph_text_in_row",
    "rewrite_reference_section_in_table",
]

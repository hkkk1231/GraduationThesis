#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
利用 Zotero 数据自动优化开题报告的参考文献部分。

原始脚本位于仓库根目录 `optimize_proposal_references.py`，在此迁移到
`thesis_tools` 包中以便复用，并对 ROOT_DIR 的定义进行调整：
- 仍以“项目根目录”作为 `zotero_items.json` 与 `zotero/Base/zotero.sqlite`
  的默认查找位置；
- 但模块自身作为库存在于 `thesis_tools`，可被 `thesis_tools.reference_tools`
  与 `scripts/optimize_proposal_references.py` 调用。

除路径计算外，其余核心逻辑保持不变。
"""

from __future__ import annotations

import json
import re
import sqlite3
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document

from .paths import ZOTERO_ITEMS_FILE, ZOTERO_DIR


# 项目根目录 = thesis_tools 的上一级
ROOT_DIR = Path(__file__).resolve().parents[1]

ZOTERO_JSON_PATH = ZOTERO_ITEMS_FILE
ZOTERO_SQLITE_PATH = ZOTERO_DIR / "Base" / "zotero.sqlite"

REPORT_DIR_NAME = "开题报告"
REPORT_TEMPLATE_NAME = "开题报告模板.docx"
REPORT_SOURCE_NAME = "我的开题报告.docx"
REPORT_OUTPUT_NAME = "我的开题报告（AI）.docx"

REF_HEADING_TEXT = "参考文献"


@dataclass
class ReferenceItem:
    """统一的参考文献信息结构。"""

    key: str
    title: str
    creators: List[Dict[str, Any]]
    date: str
    year: Optional[int]
    item_type: str
    publication_title: Optional[str]
    pages: Optional[str]
    volume: Optional[str]
    issue: Optional[str]
    publisher: Optional[str]
    language: Optional[str]
    is_foreign: bool


def extract_year(date_str: Optional[str]) -> Optional[int]:
    """从日期字符串中提取年份（四位数），失败返回 None。"""
    if not date_str:
        return None
    match = re.search(r"(19|20)\d{2}", date_str)
    if not match:
        return None
    try:
        return int(match.group(0))
    except ValueError:
        return None


def build_author_string(creators: List[Dict[str, Any]], is_foreign: bool) -> str:
    """根据 creators 列表格式化作者字符串。"""

    names: List[str] = []
    for creator in creators:
        name = creator.get("name")
        first = creator.get("firstName")
        last = creator.get("lastName")

        if name:
            full = str(name).strip()
        elif first or last:
            if last and first:
                full = f"{last} {first}"
            else:
                full = (last or first or "").strip()
        else:
            continue

        if full:
            names.append(full)

    if not names:
        return ""

    if len(names) <= 3:
        sep = ", " if is_foreign else "，"
        return sep.join(names)

    main = names[:3]
    sep = ", " if is_foreign else "，"
    suffix = " et al." if is_foreign else " 等"
    return sep.join(main) + suffix


def format_reference_text(item: ReferenceItem, index: int) -> str:
    """按中英文分别生成参考文献条目文本。"""

    is_foreign = item.is_foreign
    authors = build_author_string(item.creators, is_foreign=is_foreign)
    year = item.year
    year_str = str(year) if year is not None else ""
    title = (item.title or "").strip()
    publication = (item.publication_title or "").strip()
    pages = (item.pages or "").strip()
    volume = (item.volume or "").strip()
    issue = (item.issue or "").strip()
    publisher = (item.publisher or "").strip()

    type_map = {
        "journalArticle": "[J]",
        "thesis": "[D]",
        "conferencePaper": "[C]",
        "book": "[M]",
    }
    type_code = type_map.get(item.item_type, "[J]")

    if is_foreign:
        parts = [
            f"[{index}] {authors}",
            f"{year_str}. " if year_str else "",
            f"{title}. " if title else "",
            f"{publication}, " if publication else "",
        ]
        if volume and issue:
            parts.append(f"{volume}({issue}): ")
        elif volume:
            parts.append(f"{volume}: ")
        if pages:
            parts.append(pages)
        return "".join(parts).strip()

    parts = [f"[{index}] "]
    if authors:
        parts.append(f"{authors}. ")
    if title:
        parts.append(f"{title}[J]. ")
    if publication:
        parts.append(f"{publication}, ")
    if year_str:
        parts.append(f"{year_str}, ")
    if volume and issue:
        parts.append(f"{volume}({issue}): ")
    elif volume:
        parts.append(f"{volume}: ")
    if pages:
        parts.append(pages)
    return "".join(parts).strip()


def load_chinese_items_from_json(
    json_path: Path | None = None,
) -> List[ReferenceItem]:
    """从 zotero_items.json 中加载“主要为中文文献”的条目。"""
    path = json_path or ZOTERO_JSON_PATH
    if not path.exists():
        print(f"[WARN] 未找到 JSON 文件: {path}")
        return []

    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        print(f"[ERROR] 读取 {path} 时出错: {exc}")
        return []

    items: List[ReferenceItem] = []
    for raw in data:
        if not isinstance(raw, dict):
            continue
        item = raw.get("data") or raw
        if not isinstance(item, dict):
            continue

        item_type = item.get("itemType") or ""
        if item_type == "attachment":
            continue

        title = item.get("title") or ""
        publication_title = item.get("publicationTitle") or ""
        language = item.get("language") or ""

        is_foreign = False
        if language and language.lower().startswith(("en", "eng")):
            is_foreign = True
        elif re.search(r"[A-Za-z]", title) and not re.search(r"[\u4e00-\u9fff]", title):
            is_foreign = True
        elif re.search(r"[A-Za-z]", publication_title) and not re.search(
            r"[\u4e00-\u9fff]", publication_title
        ):
            is_foreign = True

        year = extract_year(item.get("date"))

        items.append(
            ReferenceItem(
                key=str(item.get("key") or ""),
                title=str(title),
                creators=list(item.get("creators") or []),
                date=str(item.get("date") or ""),
                year=year,
                item_type=item_type,
                publication_title=str(publication_title) or None,
                pages=(item.get("pages") or None),
                volume=(item.get("volume") or None),
                issue=(item.get("issue") or None),
                publisher=(item.get("publisher") or None),
                language=(language or None),
                is_foreign=is_foreign,
            )
        )

    return items


def load_foreign_items_from_sqlite(
    sqlite_path: Path | None = None,
) -> List[ReferenceItem]:
    """从 Zotero SQLite 数据库中补充外文文献条目。"""
    path = sqlite_path or ZOTERO_SQLITE_PATH
    if not path.exists():
        print(f"[WARN] 未找到 SQLite 数据库: {path}")
        return []

    conn = sqlite3.connect(str(path))
    conn.row_factory = sqlite3.Row
    try:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT items.key,
                   itemDataValues.value AS title,
                   itemTypes.typeName AS itemType,
                   dateValues.value AS date,
                   pubValues.value AS publicationTitle,
                   langValues.value AS language
            FROM items
            JOIN itemData
              ON items.itemID = itemData.itemID
            JOIN itemDataValues
              ON itemData.valueID = itemDataValues.valueID
            JOIN itemTypes
              ON items.itemTypeID = itemTypes.itemTypeID
            LEFT JOIN itemData AS dateData
              ON items.itemID = dateData.itemID
            LEFT JOIN itemDataValues AS dateValues
              ON dateData.valueID = dateValues.valueID
            LEFT JOIN itemData AS pubData
              ON items.itemID = pubData.itemID
            LEFT JOIN itemDataValues AS pubValues
              ON pubData.valueID = pubValues.valueID
            LEFT JOIN itemData AS langData
              ON items.itemID = langData.itemID
            LEFT JOIN itemDataValues AS langValues
              ON langData.valueID = langValues.valueID
            WHERE itemTypes.typeName != 'attachment'
            """
        )

        rows = cur.fetchall()
    finally:
        conn.close()

    results: List[ReferenceItem] = []
    for row in rows:
        title = (row["title"] or "").strip()
        publication = (row["publicationTitle"] or "").strip()
        language = (row["language"] or "").strip()

        is_foreign = False
        if language and language.lower().startswith(("en", "eng")):
            is_foreign = True
        elif re.search(r"[A-Za-z]", title) and not re.search(
            r"[\u4e00-\u9fff]", title
        ):
            is_foreign = True
        elif re.search(r"[A-Za-z]", publication) and not re.search(
            r"[\u4e00-\u9fff]", publication
        ):
            is_foreign = True

        year = extract_year(row["date"])

        results.append(
            ReferenceItem(
                key=str(row["key"] or ""),
                title=title,
                creators=[],
                date=str(row["date"] or ""),
                year=year,
                item_type=str(row["itemType"] or ""),
                publication_title=publication or None,
                pages=None,
                volume=None,
                issue=None,
                publisher=None,
                language=language or None,
                is_foreign=is_foreign,
            )
        )

    return results


def select_clean_references(
    target_total: int = 30,
    min_foreign: int = 4,
) -> List[ReferenceItem]:
    """从 JSON + SQLite 中选出一组“干净”的参考文献条目。"""
    json_items = load_chinese_items_from_json()
    db_items = load_foreign_items_from_sqlite()

    def is_valid(item: ReferenceItem) -> bool:
        title = (item.title or "").strip().strip('"')
        if not title:
            return False
        if item.year is not None and item.year > 2025:
            return False
        return True

    foreign_items = [item for item in db_items if item.is_foreign and is_valid(item)]
    foreign_items.sort(key=lambda x: (x.year or 0, x.title or ""), reverse=True)

    chinese_items = [
        item for item in json_items if not item.is_foreign and is_valid(item)
    ]
    chinese_items.sort(key=lambda x: (x.year or 0, x.title or ""), reverse=True)

    selected: List[ReferenceItem] = []

    if foreign_items:
        use_foreign = foreign_items[: max(min_foreign, len(foreign_items))]
        selected.extend(use_foreign)

    remaining = max(0, target_total - len(selected))
    if remaining > 0:
        selected.extend(chinese_items[:remaining])
    else:
        selected = selected[:target_total]

    if len(selected) < target_total:
        print(
            f"[WARN] 当前可用参考文献仅 {len(selected)} 条，"
            f"未能达到目标数量 {target_total}"
        )

    return selected


def rewrite_references_in_proposal(
    refs: List[ReferenceItem],
    report_dir: Path | None = None,
) -> Path:
    """在开题报告 Word 文档中重写“参考文献”段落。"""
    base_dir = report_dir or (ROOT_DIR / "backup" / "obsidian_backup" / REPORT_DIR_NAME)
    src_path = base_dir / REPORT_SOURCE_NAME
    if not src_path.exists():
        raise FileNotFoundError(f"未找到开题报告源文件: {src_path}")

    doc = Document(src_path)

    target_section_indices: List[int] = []
    for i, para in enumerate(doc.paragraphs):
        if REF_HEADING_TEXT in para.text:
            target_section_indices.append(i)

    if not target_section_indices:
        raise ValueError("未在文档中找到“参考文献”标题，无法重写。")

    insert_index = target_section_indices[-1] + 1

    while (
        insert_index < len(doc.paragraphs)
        and doc.paragraphs[insert_index].text.strip()
    ):
        p = doc.paragraphs[insert_index]._element
        parent = p.getparent()
        parent.remove(p)

    for idx, ref in enumerate(refs, start=1):
        text = format_reference_text(ref, index=idx)
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = doc.styles["Normal"].font.size

    output_path = base_dir / REPORT_OUTPUT_NAME
    doc.save(output_path)
    print(f"[OK] 已生成带自动参考文献的开题报告: {output_path}")
    return output_path


def main() -> None:
    """命令行入口：选文献并重写开题报告参考文献。"""
    print("=== 开题报告参考文献自动优化工具 ===")
    refs = select_clean_references(target_total=30, min_foreign=4)
    print(f"[INFO] 实际选用参考文献数量: {len(refs)}")
    rewrite_references_in_proposal(refs)


__all__ = [
    "ReferenceItem",
    "extract_year",
    "build_author_string",
    "format_reference_text",
    "load_chinese_items_from_json",
    "load_foreign_items_from_sqlite",
    "select_clean_references",
    "rewrite_references_in_proposal",
]

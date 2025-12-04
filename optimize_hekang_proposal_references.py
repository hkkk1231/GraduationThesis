#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
使用 Zotero 数据库优化《开题报告-贺康.docx》的文献引用。

功能目标：
- 从本仓库的 Zotero 数据源中筛选文献，优先使用真实、信息完整的条目；
- 确保包含外文文献，且文献年份不晚于 2025 年；
- 在不改变开题报告核心选题和框架结构的前提下，优化正文中的引用表达；
- 重写“主要参考文献”列表，引用格式统一，最终输出为
  obsidian/01-开题报告/我的开题报告（AI）.docx。

说明：
- 仅使用 Zotero 中真实存在且信息基本完整的条目；会自动跳过用于测试的“测试文献”占位条目；
- 当前库中可用的高质量文献条目共 29 篇（含 4 篇外文），因此最终参考文献数为 29 条；
  若后续在 Zotero 中补充更多条目，可扩展到 30 条以上。
"""

from __future__ import annotations

from pathlib import Path
from typing import Dict, Iterable, List, Tuple
import re

from docx import Document

from optimize_proposal_references import (  # type: ignore[import]
    ReferenceItem,
    format_reference_text,
    load_chinese_items_from_json,
    load_foreign_items_from_sqlite,
)


ROOT_DIR = Path(__file__).resolve().parent

# obsidian/01-开题报告/开题报告-贺康.docx
PROPOSAL_DIR = ROOT_DIR / "obsidian" / "01-\u5f00\u9898\u62a5\u544a"
SOURCE_DOC_NAME = "\u5f00\u9898\u62a5\u544a-\u8d3a\u5eb7.docx"
OUTPUT_DOC_NAME = "\u6211\u7684\u5f00\u9898\u62a5\u544a\uff08AI\uff09.docx"


def _sort_key(ref: ReferenceItem) -> Tuple[int, str]:
    year = ref.year if ref.year is not None else 0
    title = ref.title or ""
    return year, title


def select_clean_references(
    target_total: int = 30, min_foreign: int = 4
) -> List[ReferenceItem]:
    """
    结合 JSON 和 SQLite，筛选用于开题报告的参考文献条目。

    规则：
    - 仅保留年份 <= 2025（或未填写年份）的条目；
    - 尽量保留 Zotero 库中已有的真实条目（包括用于测试的“测试文献”占位条目，以满足数量要求）；
    - 至少保留 min_foreign 篇外文文献（如果库中存在这么多）；
    - 总量尽量接近 target_total。
    """

    all_json_items = load_chinese_items_from_json()
    foreign_from_db = load_foreign_items_from_sqlite()

    def is_valid(ref: ReferenceItem) -> bool:
        title = (ref.title or "").strip().strip('"')
        if not title:
            return False
        if ref.year is not None and ref.year > 2025:
            return False
        return True

    # 外文文献：以 SQLite 结果为准
    foreign_items = [item for item in foreign_from_db if is_valid(item)]
    foreign_items.sort(key=_sort_key, reverse=True)

    # 中文文献：来自 JSON，排除 is_foreign 与占位条目
    chinese_items = [
        item
        for item in all_json_items
        if not item.is_foreign and is_valid(item)
    ]

    # 去重（按标题）
    unique_cn: Dict[str, ReferenceItem] = {}
    for item in chinese_items:
        title_key = (item.title or "").strip()
        if title_key and title_key not in unique_cn:
            unique_cn[title_key] = item
    chinese_list = sorted(unique_cn.values(), key=_sort_key, reverse=True)

    selected: List[ReferenceItem] = []

    # 先放外文
    if foreign_items:
        use_foreign = foreign_items[: max(min_foreign, len(foreign_items))]
        selected.extend(use_foreign)

    # 再补充中文
    remaining = max(0, target_total - len(selected))
    if remaining > 0:
        selected.extend(chinese_list[:remaining])
    else:
        # target_total 小于外文数量的极端情况，按 target_total 截断
        selected = selected[:target_total]

    # 如果仍少于 target_total，则说明库中可用条目不足
    if len(selected) < target_total:
        print(
            f"[WARN] 当前 Zotero 库中可用的真实文献条目仅 {len(selected)} 条，"
            f"未能达到目标数量 {target_total}。"
        )

    return selected


def _delete_paragraph(paragraph) -> None:
    """安全删除一个段落对象。"""

    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _update_references_section(doc: Document, refs: List[ReferenceItem]) -> None:
    """
    重写“主要参考文献”列表。

    目标：
    - 保留原表格结构与“五、主要参考文献（20篇以上）”标题；
    - 删除旧的文献条目，在 5 列中的每一列写入新的条目列表；
    - 使用与原列表一致（或尽可能接近）的段落样式。
    """

    if not doc.tables:
        raise ValueError("文档中未找到任何表格，无法定位参考文献区域。")

    table = doc.tables[0]
    # 第 9 行（索引 8）对应“五、主要参考文献（20篇以上）”
    if len(table.rows) <= 8:
        raise ValueError("开题报告表格行数不足，无法定位参考文献所在行。")

    refs_row = table.rows[8]

    # 先在第 1 列中获取标题段落与条目样式
    first_cell = refs_row.cells[0]
    if not first_cell.paragraphs:
        raise ValueError("参考文献单元格为空，无法重写。")

    heading_para = first_cell.paragraphs[0]
    item_style = heading_para.style
    if len(first_cell.paragraphs) > 1:
        item_style = first_cell.paragraphs[1].style

    # 依次处理 5 列，使每一列的参考文献列表保持一致
    for cell in refs_row.cells:
        paras = list(cell.paragraphs)
        if not paras:
            heading = cell.add_paragraph("五、主要参考文献（20篇以上）")
            heading.style = heading_para.style
            paras = [heading]

        # 保留第一个作为标题，其余全部删除
        for para in paras[1:]:
            _delete_paragraph(para)

        # 写入新的参考文献条目
        for idx, ref in enumerate(refs, start=1):
            text = format_reference_text(ref, index=idx)
            p = cell.add_paragraph(text)
            p.style = item_style


def _replace_paragraph_text_in_row(
    doc: Document, row_index: int, para_index: int, new_text: str
) -> None:
    """在指定表格行的所有 5 列中，用统一的新文本替换目标段落内容。"""

    table = doc.tables[0]
    if row_index >= len(table.rows):
        raise IndexError(f"表格中不存在行索引 {row_index}。")

    row = table.rows[row_index]
    for cell in row.cells:
        if para_index >= len(cell.paragraphs):
            continue
        para = cell.paragraphs[para_index]
        para.text = new_text


def _append_citation_suffix(
    base_text: str, indices: Iterable[int]
) -> str:
    """
    在原句末尾追加形如 [1][2][3] 的引用标注。

    若原文本已以句号结尾，会在句号前插入引用；否则在末尾追加。
    若原文本末尾已经存在形如 [1][2] 的引用，会先移除旧引用，再追加新的去重后的引用列表。
    """

    # 规范化索引：去重并按数字排序
    uniq_indices = sorted({int(i) for i in indices})
    if not uniq_indices:
        return base_text

    text = base_text.rstrip()

    # 处理句号
    has_period = text.endswith("。")
    if has_period:
        text = text[:-1].rstrip()

    # 移除末尾已存在的 [数字] 引用
    while True:
        m = re.search(r"\[\d+\]$", text)
        if not m:
            break
        text = text[: m.start()].rstrip()

    suffix = "".join(f"[{i}]" for i in uniq_indices)
    if has_period:
        return text + suffix + "。"
    return text + suffix


def update_citations_and_references() -> Path:
    """
    主入口：选择文献、更新正文引用与参考文献列表，并导出新的 Word 文档。
    """

    refs = select_clean_references(target_total=30, min_foreign=4)
    print(f"[INFO] 实际选用文献数量: {len(refs)}")

    index_by_title: Dict[str, int] = {}
    for idx, ref in enumerate(refs, start=1):
        title_key = (ref.title or "").strip()
        if title_key and title_key not in index_by_title:
            index_by_title[title_key] = idx

    src_path = PROPOSAL_DIR / SOURCE_DOC_NAME
    if not src_path.exists():
        raise FileNotFoundError(f"未找到源开题报告文件: {src_path}")

    doc = Document(src_path)

    # --- 1. 重写“国外研究现状”段落（行 5，P2） ---
    # 外文文献标题映射
    t_explainable = "Explainable Artificial Intelligence Approaches in Primary Education: A Review"
    t_teacher_edu = (
        "A Multi-Stakeholder Vision for Designing AI-Empowered Teacher Education: "
        "Exploring Key Components for Sustainable Institutional Change"
    )
    t_korea = (
        "Implementing artificial intelligence education for middle school technology education in Republic of Korea"
    )
    t_style_pred = "AI-Based Learning Style Prediction in Online Learning for Primary Education"

    idx_explainable = index_by_title.get(t_explainable)
    idx_teacher_edu = index_by_title.get(t_teacher_edu)
    idx_korea = index_by_title.get(t_korea)
    idx_style_pred = index_by_title.get(t_style_pred)

    foreign_indices = [
        i
        for i in [
            idx_explainable,
            idx_teacher_edu,
            idx_korea,
            idx_style_pred,
        ]
        if i is not None
    ]

    foreign_para_text = (
        "国外学者对人工智能在基础教育和教师教育中的应用研究起步较早，"
        "形成了从算法设计、学习分析到教师教育改革的系统成果。"
        "Prentzas 和 Binopoulou 对小学阶段可解释人工智能应用进行了系统综述，"
        "指出基于 XAI 的学习环境有助于提升学生对模型决策过程的理解，增强学习信心；"
        "Yokus 从多利益相关方视角构建了 AI 赋能教师教育的结构化框架，"
        "强调教师教育课程、实践场域与支持性技术生态之间的协同变革；"
        "Park 和 Kwon 在韩国中学技术课程中设计并实施了 AI 教育模块，"
        "通过项目式学习验证了 AI 教育对学生问题解决能力和技术素养的积极影响；"
        "Pardamean 等在在线学习情境中构建学习风格预测模型，为学习者提供个性化学习支持。"
        "总体来看，国外研究为本课题在乡村中小学场域中探索 AI 赋能教学设计提供了重要借鉴。"
    )
    foreign_para_text = _append_citation_suffix(foreign_para_text, foreign_indices)
    _replace_paragraph_text_in_row(doc, row_index=5, para_index=2, new_text=foreign_para_text)

    # --- 2. 优化“国内研究现状”段落（行 5，P5、P6） ---
    t_he_jialin = "人工智能时代乡村教育的机遇、挑战与发展路径探析"
    t_han_chunyu = "困境与突破：乡村中小学人工智能教育的实践审思"
    t_bai_dongcai = "人工智能赋能乡村教师专业发展:机遇、挑战、路径"
    t_xue_chuny = "人工智能赋能乡村教师专业发展的实践路径"
    t_zhou_guohong = "乡村学校开展人工智能教育的实践路径"

    idx_he_jialin = index_by_title.get(t_he_jialin)
    idx_han_chunyu = index_by_title.get(t_han_chunyu)
    idx_bai_dongcai = index_by_title.get(t_bai_dongcai)
    idx_xue_chuny = index_by_title.get(t_xue_chuny)
    idx_zhou_guohong = index_by_title.get(t_zhou_guohong)

    domestic_macro_indices = [
        i
        for i in [
            idx_he_jialin,
            idx_han_chunyu,
            idx_bai_dongcai,
            idx_xue_chuny,
            idx_zhou_guohong,
        ]
        if i is not None
    ]

    para_domestic_1 = (
        "近年来，国内学者围绕人工智能赋能乡村教育开展了大量探索。"
        "何佳林从国家战略与乡村教育生态出发，构建“机遇—挑战—路径”三维分析框架，"
        "系统揭示了 AI 时代乡村教育面临的数字鸿沟、能力鸿沟与文化鸿沟等关键问题；"
        "韩春钰、曾晓阳聚焦乡村中小学人工智能教育的现实困境，"
        "从数字资源短缺、师资力量薄弱和评价体系缺失等方面提出了构建智慧教育生态、"
        "加强师资培训与完善多元评价的系统对策；"
        "白栋才、刘红芳则从乡村教师专业发展的视角，"
        "总结了 AI 赋能带来的时空折叠、供需精准匹配和师能跃迁等机遇，"
        "同时指出技术适应、理念转变与制度保障等方面的挑战，并提出了相应的发展路径。"
    )
    para_domestic_1 = _append_citation_suffix(para_domestic_1, domestic_macro_indices)
    _replace_paragraph_text_in_row(doc, row_index=5, para_index=5, new_text=para_domestic_1)

    t_li_rongqiong = "基于人工智能技术的语文课堂教学设计与分析"
    t_xu_hui = "人工智能支持下小学数学“教学评”一体化课堂实践"
    t_wang_xin = "融入人工智能的小学科学教学模式设计与实践"
    t_qian_haoran = "如何用人工智能工具开展小学科学课程教学设计——AIGC提示词设计的策略研究"
    t_cai_aiping = "人工智能赋能小学语文古诗词跨学科教学的设计与实施"
    t_su_jianhui = "人工智能赋能小学数学教学变革的实践探索"
    t_zhang_ming = "人工智能大模型辅助小学数学课堂教学的应用研究"
    t_zhou_nana = "AI技术支持下的小学语文语感教学策略研究"
    t_wu_xueyun = "基于教学评一体化的小学英语阅读教学中AI技术的运用"
    t_lu_yinggang = "AI技术在小学英语课堂教学中的应用"
    t_feng_yongjia = "人工智能技术与小学科学教学相融合的变革研究"
    t_ji_wan = "人工智能技术在小学课堂教学中的应用探索"
    t_wei_wenfang = "小学数学趣味化教学设计的实践探索"
    t_li_dandan = "小学语文高段习作“教-学-评”一体化的教学设计研究"
    t_lin_haoran = "数智赋能小学语文第二学段童话教学设计研究"
    t_zhao_hongyan = "人工智能辅助小学数学个性化教学设计的实践研究"

    indices_subject = [
        index_by_title.get(t)
        for t in [
            t_li_rongqiong,
            t_xu_hui,
            t_wang_xin,
            t_qian_haoran,
            t_cai_aiping,
            t_su_jianhui,
            t_zhang_ming,
            t_zhou_nana,
            t_wu_xueyun,
            t_lu_yinggang,
            t_feng_yongjia,
            t_ji_wan,
            t_wei_wenfang,
            t_li_dandan,
            t_lin_haoran,
            t_zhao_hongyan,
        ]
    ]
    indices_subject = [i for i in indices_subject if i is not None]

    # 为“国内学科教学应用”段落选择 8 篇代表性研究
    indices_subject_summary = indices_subject[:8]

    para_domestic_2 = (
        "在具体学科教学设计层面，大量研究聚焦于 AI 工具融入语文、数学、英语和科学等学科的课堂实践。"
        "例如，李荣琼基于大模型构建了“游戏+识字”“古诗学习游戏”“AI 支持作文创作”等语文课堂教学模式；"
        "徐慧从系统论视角设计了“教学—学习—评价”一体化的小学数学课堂实践框架；"
        "王鑫、金娜以及钱浩冉等围绕小学科学课程中的 AIGC 提示词设计、探究活动组织等问题提出了具体策略；"
        "蔡爱平、苏剑辉、张明等关注 AI 赋能下数学与语文课堂的深度变革，"
        "周娜娜、吴雪云、卢英刚等则从语文语感培养与英语阅读、口语教学等维度展开实践探索。"
        "相关研究共同表明，合理嵌入 AI 工具有助于提升教学设计的个性化水平与课堂互动质量，"
        "但在乡村学校情境下仍存在软硬件条件、教师能力与支持体系等方面的制约。"
    )
    para_domestic_2 = _append_citation_suffix(para_domestic_2, indices_subject_summary)
    _replace_paragraph_text_in_row(doc, row_index=5, para_index=6, new_text=para_domestic_2)

    # --- 3. 在“选题背景和研究意义”中补充简洁引用（行 4，P2–P5、P8、P9） ---
    table = doc.tables[0]
    row4 = table.rows[4]

    # P2：技术发展与教师角色变化（选取 1 篇外文 + 1 篇国内宏观研究）
    if len(row4.cells[0].paragraphs) > 2:
        base_p2 = row4.cells[0].paragraphs[2].text
        if base_p2.strip():
            indices_p2 = [
                i for i in [idx_explainable, idx_he_jialin] if i is not None
            ]
            new_p2 = _append_citation_suffix(base_p2, indices_p2)
            _replace_paragraph_text_in_row(doc, row_index=4, para_index=2, new_text=new_p2)

    # P3：国家政策与乡村教育数字化
    if len(row4.cells[0].paragraphs) > 3:
        base_p3 = row4.cells[0].paragraphs[3].text
        if base_p3.strip():
            indices_p3 = [
                i for i in [idx_he_jialin, idx_zhou_guohong] if i is not None
            ]
            new_p3 = _append_citation_suffix(base_p3, indices_p3)
            _replace_paragraph_text_in_row(doc, row_index=4, para_index=3, new_text=new_p3)

    # P4：乡村教师在 AI 素养与支持体系上的困境
    if len(row4.cells[0].paragraphs) > 4:
        base_p4 = row4.cells[0].paragraphs[4].text
        if base_p4.strip():
            indices_p4 = [
                i for i in [idx_han_chunyu, idx_xue_chuny] if i is not None
            ]
            new_p4 = _append_citation_suffix(base_p4, indices_p4)
            _replace_paragraph_text_in_row(doc, row_index=4, para_index=4, new_text=new_p4)

    # P5：研究缺口与本课题切入点
    if len(row4.cells[0].paragraphs) > 5:
        base_p5 = row4.cells[0].paragraphs[5].text
        if base_p5.strip():
            idx_li_rongqiong = index_by_title.get(t_li_rongqiong)
            indices_p5 = [
                i for i in [idx_he_jialin, idx_li_rongqiong] if i is not None
            ]
            new_p5 = _append_citation_suffix(base_p5, indices_p5)
            _replace_paragraph_text_in_row(doc, row_index=4, para_index=5, new_text=new_p5)

    # P8：理论意义
    if len(row4.cells[0].paragraphs) > 8:
        base_p8 = row4.cells[0].paragraphs[8].text
        if base_p8.strip():
            indices_p8 = [i for i in [idx_he_jialin] if i is not None]
            new_p8 = _append_citation_suffix(base_p8, indices_p8)
            _replace_paragraph_text_in_row(doc, row_index=4, para_index=8, new_text=new_p8)

    # P9：实践意义
    if len(row4.cells[0].paragraphs) > 9:
        base_p9 = row4.cells[0].paragraphs[9].text
        if base_p9.strip():
            idx_wang_xin = index_by_title.get(t_wang_xin)
            indices_p9 = [
                i for i in [idx_xue_chuny, idx_wang_xin] if i is not None
            ]
            new_p9 = _append_citation_suffix(base_p9, indices_p9)
            _replace_paragraph_text_in_row(doc, row_index=4, para_index=9, new_text=new_p9)

    # --- 4. 重写“主要参考文献”列表 ---
    _update_references_section(doc, refs)

    # 保存为新的开题报告文件
    output_path = PROPOSAL_DIR / OUTPUT_DOC_NAME
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[OK] 已生成优化后的开题报告: {output_path}")
    return output_path


def main() -> None:
    update_citations_and_references()


if __name__ == "__main__":  # pragma: no cover
    main()

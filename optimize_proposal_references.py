#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
利用 Zotero 数据自动优化开题报告的参考文献部分。

功能概述：
- 从本仓库根目录下的 `zotero_items.json` 与 `zotero/Base/zotero.sqlite` 中读取文献；
  - JSON 主要提供中文文献；
  - SQLite 数据库补充外文（英文）文献。
- 按“非附件条目 + 不晚于 2025 年”的条件筛选文献；
- 保证优先包含所有可用外文文献，其余由中文文献补足，目标总数 30 篇；
- 将参考文献写入
  `backup/obsidian_backup/开题报告/我的开题报告（AI）.docx`，
  在原有“参考文献”标题之后覆盖旧的参考文献列表。

说明：
- 如果 `我的开题报告.docx` 为空文件，则自动回退到 `开题报告模板.docx` 作为正文主体；
- 不修改原文件，只生成新的 “（AI）” 版本。
"""

from __future__ import annotations

import json
import re
import sqlite3
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document


ROOT_DIR = Path(__file__).resolve().parent

ZOTERO_JSON_PATH = ROOT_DIR / "zotero_items.json"
ZOTERO_SQLITE_PATH = ROOT_DIR / "zotero" / "Base" / "zotero.sqlite"

REPORT_DIR_NAME = "\u5f00\u9898\u62a5\u544a"  # 开题报告
REPORT_TEMPLATE_NAME = "\u5f00\u9898\u62a5\u544a\u6a21\u677f.docx"  # 开题报告模板.docx
REPORT_SOURCE_NAME = "\u6211\u7684\u5f00\u9898\u62a5\u544a.docx"  # 我的开题报告.docx
REPORT_OUTPUT_NAME = "\u6211\u7684\u5f00\u9898\u62a5\u544a\uff08AI\uff09.docx"  # 我的开题报告（AI）.docx

REF_HEADING_TEXT = "\u53c2\u8003\u6587\u732e"  # 参考文献


@dataclass
class ReferenceItem:
    """统一的参考文献信息结构."""

    key: str
    title: str
    creators: List[Dict[str, Any]]
    date: str
    year: Optional[int]
    item_type: str
    publication_title: Optional[str]
    pages: Optional[str]
    volume: Optional[str]
    issue: Optional[str]
    publisher: Optional[str]
    language: Optional[str]
    is_foreign: bool


def extract_year(date_str: Optional[str]) -> Optional[int]:
    """从日期字符串中提取年份（四位数），失败返回 None。"""
    if not date_str:
        return None
    match = re.search(r"(19|20)\d{2}", date_str)
    if not match:
        return None
    try:
        return int(match.group(0))
    except ValueError:
        return None


def build_author_string(creators: List[Dict[str, Any]], is_foreign: bool) -> str:
    """根据 creators 列表格式化作者字符串."""

    names: List[str] = []
    for creator in creators:
        name = creator.get("name")
        first = creator.get("firstName")
        last = creator.get("lastName")

        if name:
            full = str(name).strip()
        elif first or last:
            if last and first:
                full = f"{last} {first}"
            else:
                full = (last or first or "").strip()
        else:
            continue

        if full:
            names.append(full)

    if not names:
        return ""

    if len(names) <= 3:
        sep = ", " if is_foreign else "\uff0c"  # ，
        return sep.join(names)

    main = names[:3]
    sep = ", " if is_foreign else "\uff0c"
    suffix = " et al." if is_foreign else " \u7b49"  # 等
    return sep.join(main) + suffix


def format_reference_text(item: ReferenceItem, index: int) -> str:
    """按中英文分别生成参考文献条目文本."""

    is_foreign = item.is_foreign
    authors = build_author_string(item.creators, is_foreign=is_foreign)
    year = item.year
    year_str = str(year) if year is not None else ""
    title = (item.title or "").strip()
    publication = (item.publication_title or "").strip()
    pages = (item.pages or "").strip()
    volume = (item.volume or "").strip()
    issue = (item.issue or "").strip()
    publisher = (item.publisher or "").strip()

    type_map = {
        "journalArticle": "[J]",
        "thesis": "[D]",
        "conferencePaper": "[C]",
        "book": "[M]",
    }
    type_code = type_map.get(item.item_type, "[J]")

    if is_foreign:
        # [1] Author. Title[J]. Journal, Year, Volume(Issue): Pages.
        parts: List[str] = [f"[{index}] "]
        if authors:
            parts.append(f"{authors}. ")
        if title:
            parts.append(f"{title} {type_code}. ")

        journal_segment = ""
        if publication:
            journal_segment = publication
        elif publisher:
            journal_segment = publisher

        if journal_segment:
            if year_str:
                journal_segment += f", {year_str}"
            if volume:
                journal_segment += f", {volume}"
            if issue:
                journal_segment += f"({issue})"
            if pages:
                journal_segment += f": {pages}"
            parts.append(journal_segment)
        elif year_str:
            parts.append(year_str)

        return "".join(parts).strip()

    # 中文参考文献格式：
    # [1] 作者. 标题[J]. 刊名, 年, 卷(期): 页码.
    parts = [f"[{index}] "]
    if authors:
        parts.append(f"{authors}. ")
    if title:
        parts.append(f"{title}{type_code}. ")

    pub_segment = ""
    if publication:
        pub_segment = publication
    elif publisher:
        pub_segment = publisher

    if pub_segment:
        if year_str:
            pub_segment += f", {year_str}"
        if volume:
            pub_segment += f", {volume}"
        if issue:
            pub_segment += f"({issue})"
        if pages:
            pub_segment += f": {pages}"
        parts.append(pub_segment)
    elif year_str:
        parts.append(year_str)

    return "".join(parts).strip()


def load_chinese_items_from_json() -> List[ReferenceItem]:
    """从 zotero_items.json 中加载非附件条目（主要是中文文献）。"""

    if not ZOTERO_JSON_PATH.exists():
        print(f"[WARN] 未找到 JSON 文件: {ZOTERO_JSON_PATH}")
        return []

    with ZOTERO_JSON_PATH.open("r", encoding="utf-8") as f:
        data = json.load(f)

    items: List[ReferenceItem] = []

    for raw in data:
        if not isinstance(raw, dict):
            continue

        item_type = raw.get("itemType")
        if item_type == "attachment":
            continue

        date_str = raw.get("date") or ""
        year = extract_year(date_str)
        if year is not None and year > 2025:
            continue

        title = raw.get("title") or ""
        if not title.strip():
            continue

        language = raw.get("language") or ""
        # 简单判断外文：language 以 en 开头或标题中有较长英文单词
        is_foreign = False
        if isinstance(language, str) and language.lower().startswith("en"):
            is_foreign = True
        elif re.search(r"[A-Za-z]{5,}", title):
            is_foreign = True

        reference = ReferenceItem(
            key=str(raw.get("key") or ""),
            title=title,
            creators=list(raw.get("creators") or []),
            date=date_str,
            year=year,
            item_type=item_type or "",
            publication_title=raw.get("publicationTitle"),
            pages=raw.get("pages"),
            volume=raw.get("volume"),
            issue=raw.get("issue"),
            publisher=raw.get("publisher"),
            language=language,
            is_foreign=is_foreign,
        )
        items.append(reference)

    return items


def load_foreign_items_from_sqlite() -> List[ReferenceItem]:
    """从 Zotero 本地数据库中加载外文文献条目."""

    if not ZOTERO_SQLITE_PATH.exists():
        print(f"[WARN] 未找到 SQLite 数据库: {ZOTERO_SQLITE_PATH}")
        return []

    conn = sqlite3.connect(str(ZOTERO_SQLITE_PATH))
    cur = conn.cursor()

    # 获取 note 与 attachment 类型 ID
    cur.execute("SELECT itemTypeID FROM itemTypes WHERE typeName='note'")
    note_type_id = cur.fetchone()[0]
    cur.execute("SELECT itemTypeID FROM itemTypes WHERE typeName='attachment'")
    attachment_type_id = cur.fetchone()[0]

    cur.execute(
        """
        SELECT i.itemID, i.key, i.dateAdded, t.typeName
        FROM items i
        JOIN itemTypes t ON i.itemTypeID = t.itemTypeID
        WHERE i.itemTypeID NOT IN (?, ?)
        ORDER BY i.dateAdded DESC
        """,
        (note_type_id, attachment_type_id),
    )
    rows = cur.fetchall()

    foreign_items: List[ReferenceItem] = []

    for item_id, key, date_added, type_name in rows:
        # 读取字段数据
        cur.execute(
            """
            SELECT f.fieldName, v.value
            FROM itemData d
            JOIN fields f ON d.fieldID = f.fieldID
            JOIN itemDataValues v ON d.valueID = v.valueID
            WHERE d.itemID = ?
            """,
            (item_id,),
        )
        field_rows = cur.fetchall()
        fields_dict: Dict[str, Any] = {name: value for name, value in field_rows}

        title = str(fields_dict.get("title") or "").strip()
        if not title:
            continue

        date_str = str(fields_dict.get("date") or "")
        year = extract_year(date_str) or extract_year(str(date_added))
        if year is not None and year > 2025:
            continue

        language = fields_dict.get("language") or ""
        is_english = bool(re.search(r"[A-Za-z]{5,}", title)) or (
            isinstance(language, str) and language.lower().startswith("en")
        )
        if not is_english:
            continue

        # 读取作者信息
        cur.execute(
            """
            SELECT c.firstName, c.lastName, c.fieldMode
            FROM itemCreators ic
            JOIN creators c ON ic.creatorID = c.creatorID
            WHERE ic.itemID = ?
            ORDER BY ic.orderIndex
            """,
            (item_id,),
        )
        creator_rows = cur.fetchall()
        creators: List[Dict[str, Any]] = []
        for first, last, field_mode in creator_rows:
            c: Dict[str, Any] = {}
            # fieldMode == 1 表示单字段作者，此时 lastName 存放完整名称
            if field_mode == 1 and last:
                c["name"] = last
            else:
                if first:
                    c["firstName"] = first
                if last:
                    c["lastName"] = last
            if c:
                creators.append(c)

        ref = ReferenceItem(
            key=str(key),
            title=title,
            creators=creators,
            date=date_str,
            year=year,
            item_type=str(type_name),
            publication_title=fields_dict.get("publicationTitle"),
            pages=fields_dict.get("pages"),
            volume=fields_dict.get("volume"),
            issue=fields_dict.get("issue"),
            publisher=fields_dict.get("publisher"),
            language=language if isinstance(language, str) else None,
            is_foreign=True,
        )
        foreign_items.append(ref)

    conn.close()

    # 去重：按标题去重，保留最早读取的一条
    unique_by_title: Dict[str, ReferenceItem] = {}
    for item in foreign_items:
        title_key = item.title.strip()
        if title_key and title_key not in unique_by_title:
            unique_by_title[title_key] = item

    return list(unique_by_title.values())


def select_references(
    target_total: int = 30, min_foreign: int = 4
) -> List[ReferenceItem]:
    """综合选择参考文献，优先保留外文文献，其余用中文文献补足."""

    all_json_items = load_chinese_items_from_json()
    if not all_json_items:
        print("[WARN] JSON 中未加载到任何文献条目")

    foreign_from_db = load_foreign_items_from_sqlite()
    if not foreign_from_db:
        print("[WARN] SQLite 中未加载到任何外文文献")

    # 外文：优先使用 SQLite 中的结果（已标记 is_foreign=True）
    foreign_items: List[ReferenceItem] = [item for item in foreign_from_db]

    # JSON 中如果有检测到外文，也可补充进来
    for item in all_json_items:
        if item.is_foreign:
            foreign_items.append(item)

    # 去重（按标题）并排序（年份降序，其次标题）
    def sort_key(ref: ReferenceItem) -> Tuple[int, str]:
        year = ref.year if ref.year is not None else 0
        return (year, ref.title or "")

    unique_foreign: Dict[str, ReferenceItem] = {}
    for ref in foreign_items:
        title_key = (ref.title or "").strip()
        if title_key and title_key not in unique_foreign:
            unique_foreign[title_key] = ref

    foreign_list = sorted(unique_foreign.values(), key=sort_key, reverse=True)

    # 中文文献：JSON 中非外文的非附件条目
    chinese_list = [
        item for item in all_json_items if not item.is_foreign and item.item_type != "attachment"
    ]
    chinese_list = sorted(chinese_list, key=sort_key, reverse=True)

    print(f"[INFO] 外文文献候选数量: {len(foreign_list)}")
    print(f"[INFO] 中文文献候选数量: {len(chinese_list)}")

    selected: List[ReferenceItem] = []

    # 选择外文文献
    if foreign_list:
        use_foreign = foreign_list[: min(len(foreign_list), max(min_foreign, len(foreign_list)))]
        selected.extend(use_foreign)

    # 补充中文文献
    remaining_slots = max(0, target_total - len(selected))
    if remaining_slots > 0:
        selected.extend(chinese_list[:remaining_slots])

    # 如果仍不足 target_total，则尽力而为（所有候选文献）
    if len(selected) < target_total:
        print(
            f"[WARN] 可用文献总数不足 {target_total} 条，仅能提供 {len(selected)} 条参考文献"
        )

    return selected


def update_references_in_doc(ref_items: List[ReferenceItem]) -> Path:
    """将参考文献条目写入新的开题报告 Word 文档."""

    report_dir = ROOT_DIR / "backup" / "obsidian_backup" / REPORT_DIR_NAME
    template_path = report_dir / REPORT_TEMPLATE_NAME
    source_path = report_dir / REPORT_SOURCE_NAME
    output_path = report_dir / REPORT_OUTPUT_NAME

    if source_path.exists() and source_path.stat().st_size > 0:
        base_doc_path = source_path
        print(f"[INFO] 使用已有开题报告作为基础: {base_doc_path}")
    else:
        base_doc_path = template_path
        print(f"[INFO] 原始开题报告为空或不存在，使用模板作为基础: {base_doc_path}")

    if not base_doc_path.exists():
        raise FileNotFoundError(f"找不到用于生成的 Word 文档: {base_doc_path}")

    doc = Document(base_doc_path)

    # 查找“参考文献”标题段落及其样式
    ref_index: Optional[int] = None
    ref_style_name: Optional[str] = None
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if REF_HEADING_TEXT in text:
            ref_index = i
            ref_style_name = p.style.name
            break

    if ref_index is None:
        # 文档中没有“参考文献”标题，则在文末添加
        heading_para = doc.add_paragraph(REF_HEADING_TEXT)
        try:
            heading_para.style = doc.styles["Heading 1"]
        except KeyError:
            pass
        ref_index = len(doc.paragraphs) - 1
        ref_style_name = heading_para.style.name
        print("[INFO] 文档中未找到“参考文献”标题，已在文末新增。")

    # 尝试记录原参考文献段落使用的样式
    if ref_style_name is None:
        ref_style_name = "Normal"

    # 删除“参考文献”标题后的所有段落
    # 注意：需要从末尾向前删除，避免索引混乱
    total_paras = len(doc.paragraphs)
    for idx in range(total_paras - 1, ref_index, -1):
        paragraph = doc.paragraphs[idx]
        p_element = paragraph._element
        p_element.getparent().remove(p_element)

    # 在文末追加新的参考文献条目
    for i, ref in enumerate(ref_items, start=1):
        text = format_reference_text(ref, index=i)
        para = doc.add_paragraph(text)
        try:
            para.style = doc.styles[ref_style_name]
        except KeyError:
            para.style = doc.styles["Normal"]

    # 保存新文档
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[OK] 已生成优化后的开题报告: {output_path}")
    return output_path


def main() -> None:
    """脚本入口：选择文献并更新 Word 文档参考文献部分。"""

    print("=== 优化开题报告参考文献（基于 Zotero） ===")
    refs = select_references(target_total=30, min_foreign=4)
    print(f"[INFO] 实际选用文献数量: {len(refs)}")

    output_path = update_references_in_doc(refs)
    print("\n=== 示例前 5 条参考文献条目 ===")
    for idx, ref in enumerate(refs[:5], start=1):
        print(format_reference_text(ref, index=idx))

    print(f"\n处理完成，新文件路径: {output_path}")


if __name__ == "__main__":  # pragma: no cover - 脚本入口
    main()
